"""
使用 python-docx 库拆分 DOCX 文档（纯Python跨平台方案）
支持 Linux/macOS/Windows
依赖: pip install python-docx
"""
import os
import sys
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls


def estimate_page_breaks(doc):
    """估算文档中的分页符位置（段落索引）"""
    page_breaks = [0]  # 第一页从索引0开始
    
    for i, para in enumerate(doc.paragraphs):
        # 检查段落是否包含分页符
        if para._element.xpath('.//w:br[@w:type="page"]', namespaces=para._element.nsmap):
            page_breaks.append(i + 1)
        
        # 检查段落格式中的分页设置
        if para.paragraph_format.page_break_before:
            if i not in page_breaks:
                page_breaks.append(i)
    
    return page_breaks


def count_pages_docx(doc):
    """估算DOCX文档的页数"""
    # 方法1: 查找所有分页符
    page_breaks = estimate_page_breaks(doc)
    estimated_pages = len(page_breaks)
    
    # 方法2: 根据内容估算（如果没有明显的分页符）
    if estimated_pages <= 1:
        total_paragraphs = len(doc.paragraphs)
        # 假设每页约20个段落（粗略估算）
        estimated_pages = max(1, (total_paragraphs + 19) // 20)
    
    return estimated_pages


def copy_paragraph(source_para, target_doc):
    """复制段落到目标文档"""
    new_para = target_doc.add_paragraph()
    
    # 复制段落格式
    if source_para.style:
        try:
            new_para.style = source_para.style
        except:
            pass
    
    # 复制段落级别格式
    try:
        new_para.paragraph_format.alignment = source_para.paragraph_format.alignment
        new_para.paragraph_format.left_indent = source_para.paragraph_format.left_indent
        new_para.paragraph_format.right_indent = source_para.paragraph_format.right_indent
        new_para.paragraph_format.first_line_indent = source_para.paragraph_format.first_line_indent
        new_para.paragraph_format.space_before = source_para.paragraph_format.space_before
        new_para.paragraph_format.space_after = source_para.paragraph_format.space_after
        new_para.paragraph_format.line_spacing = source_para.paragraph_format.line_spacing
    except:
        pass
    
    # 复制文本内容和run格式
    for run in source_para.runs:
        new_run = new_para.add_run(run.text)
        try:
            new_run.bold = run.bold
            new_run.italic = run.italic
            new_run.underline = run.underline
            if run.font.size:
                new_run.font.size = run.font.size
            if run.font.name:
                new_run.font.name = run.font.name
            if run.font.color.rgb:
                new_run.font.color.rgb = run.font.color.rgb
        except:
            pass
    
    return new_para


def copy_table(source_table, target_doc):
    """复制表格到目标文档"""
    # 获取表格的行列数
    rows_count = len(source_table.rows)
    cols_count = len(source_table.columns)
    
    # 创建新表格
    new_table = target_doc.add_table(rows=rows_count, cols=cols_count)
    
    # 尝试复制表格样式
    try:
        if source_table.style:
            new_table.style = source_table.style
    except:
        pass
    
    # 复制每个单元格的内容
    for i, row in enumerate(source_table.rows):
        for j, cell in enumerate(row.cells):
            new_cell = new_table.rows[i].cells[j]
            # 复制单元格中的段落
            new_cell.text = ''  # 清空默认段落
            for para in cell.paragraphs:
                copy_paragraph(para, new_cell)
    
    return new_table


def split_docx_by_sections(input_path: str, output_dir: str, pages_per_file: int, original_filename: str = None):
    """
    按段落和分页符拆分DOCX文档
    注意：由于python-docx无法精确获取页数，此方法按段落分组拆分
    """
    if original_filename is None:
        original_filename = Path(input_path).stem
    
    print(f"开始拆分文档: {input_path}")
    print(f"输出目录: {output_dir}")
    print(f"目标：每 {pages_per_file} 页一个文件（基于段落估算）")
    
    # 清理并创建输出目录
    import shutil
    if os.path.exists(output_dir):
        try:
            shutil.rmtree(output_dir)
            print(f"已清理旧的输出目录")
        except Exception as e:
            print(f"清理目录时出错: {e}")
    
    os.makedirs(output_dir, exist_ok=True)
    
    # 加载文档
    print("加载文档...")
    doc = Document(input_path)
    
    # 估算总页数
    total_pages = count_pages_docx(doc)
    print(f"估算总页数: {total_pages}")
    
    # 计算总文件数
    total_files = (total_pages + pages_per_file - 1) // pages_per_file
    print(f"预计生成文件数: {total_files}")
    print(f"PROGRESS:TOTAL_FILES:{total_files}")
    
    # 计算每个文件大约应包含多少段落
    total_paragraphs = len(doc.paragraphs)
    total_tables = len(doc.tables)
    
    print(f"文档总段落数: {total_paragraphs}")
    print(f"文档总表格数: {total_tables}")
    
    # 按段落数平均分配
    paras_per_file = max(1, total_paragraphs // total_files)
    
    # 开始拆分
    file_index = 1
    para_index = 0
    
    while para_index < total_paragraphs:
        print(f"\n正在创建第 {file_index} 个文件...")
        print(f"PROGRESS:FILE_START:{file_index}:{total_files}")
        
        # 创建新文档
        print(f"PROGRESS:FILE_STEP:{file_index}:创建新文档:10")
        new_doc = Document()
        
        # 复制文档的节设置
        try:
            new_section = new_doc.sections[0]
            source_section = doc.sections[0]
            new_section.page_height = source_section.page_height
            new_section.page_width = source_section.page_width
            new_section.left_margin = source_section.left_margin
            new_section.right_margin = source_section.right_margin
            new_section.top_margin = source_section.top_margin
            new_section.bottom_margin = source_section.bottom_margin
        except Exception as e:
            print(f"  复制页面设置失败: {e}")
        
        # 确定本文件的段落范围
        end_para_index = min(para_index + paras_per_file, total_paragraphs)
        
        # 如果是最后一个文件，包含所有剩余段落
        if file_index == total_files:
            end_para_index = total_paragraphs
        
        print(f"  复制段落 {para_index + 1} 到 {end_para_index}...")
        print(f"PROGRESS:FILE_STEP:{file_index}:复制内容:50")
        
        # 复制段落
        copied_count = 0
        for i in range(para_index, end_para_index):
            try:
                copy_paragraph(doc.paragraphs[i], new_doc)
                copied_count += 1
            except Exception as e:
                print(f"  复制段落 {i} 时出错: {e}")
        
        print(f"  已复制 {copied_count} 个段落")
        
        # 生成输出文件名
        start_page = (file_index - 1) * pages_per_file + 1
        end_page = min(file_index * pages_per_file, total_pages)
        
        if start_page == end_page:
            out_filename = f"{original_filename} (第{start_page}页).docx"
        else:
            out_filename = f"{original_filename} (第{start_page}-{end_page}页).docx"
        
        out_path = os.path.join(output_dir, out_filename)
        
        # 保存文件
        print(f"  保存文件: {out_filename}")
        print(f"PROGRESS:FILE_STEP:{file_index}:保存文件:90")
        new_doc.save(out_path)
        
        print(f"已保存: {out_filename}")
        print(f"PROGRESS:FILE_COMPLETE:{file_index}:{total_files}")
        
        # 更新索引
        file_index += 1
        para_index = end_para_index
    
    print(f"\n拆分完成！共生成 {file_index - 1} 个文件")
    print(f"PROGRESS:ALL_FILES_COMPLETE:{file_index - 1}:{total_files}")


def main():
    """主函数"""
    if len(sys.argv) not in [4, 5]:
        print("用法: python split_docx_pages_python_docx.py <输入文件> <输出目录> <每文件页数> [原始文件名]")
        sys.exit(1)
    
    input_path = sys.argv[1]
    output_dir = sys.argv[2]
    pages_per_file = int(sys.argv[3])
    original_filename = sys.argv[4] if len(sys.argv) == 5 else None
    
    if not os.path.exists(input_path):
        print(f"错误: 输入文件不存在: {input_path}")
        sys.exit(1)
    
    if pages_per_file < 1 or pages_per_file > 1000:
        print(f"错误: 每文件页数必须在 1-1000 之间")
        sys.exit(1)
    
    try:
        split_docx_by_sections(input_path, output_dir, pages_per_file, original_filename)
        print("拆分成功!")
    except Exception as e:
        print(f"拆分失败: {e}")
        import traceback
        traceback.print_exc()
        sys.exit(1)


if __name__ == "__main__":
    main()
