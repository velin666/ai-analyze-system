#!/usr/bin/env python3
"""
跨平台 DOCX 拆分测试脚本
测试不同平台下的拆分功能
"""
import os
import sys
import platform
import subprocess
import tempfile
from pathlib import Path


def print_section(title):
    """打印分节标题"""
    print(f"\n{'='*60}")
    print(f"  {title}")
    print(f"{'='*60}\n")


def check_dependencies():
    """检查依赖是否安装"""
    print_section("检查系统依赖")
    
    system = platform.system()
    print(f"操作系统: {system} {platform.release()}")
    print(f"Python 版本: {platform.python_version()}")
    
    # 检查 win32com (Windows)
    if system == 'Windows':
        try:
            import win32com.client
            print("✓ win32com 已安装")
            has_win32com = True
        except ImportError:
            print("✗ win32com 未安装 (pip install pywin32)")
            has_win32com = False
    else:
        has_win32com = False
    
    # 检查 LibreOffice UNO
    try:
        import uno
        print("✓ LibreOffice UNO 已安装")
        has_uno = True
    except ImportError:
        print("✗ LibreOffice UNO 未安装")
        if system == 'Linux':
            print("  安装方法: sudo apt-get install python3-uno")
        elif system == 'Darwin':
            print("  安装方法: brew install libreoffice")
        has_uno = False
    
    # 检查 LibreOffice 服务
    if has_uno:
        try:
            import socket
            sock = socket.socket(socket.AF_INET, socket.SOCK_STREAM)
            result = sock.connect_ex(('127.0.0.1', 2002))
            sock.close()
            
            if result == 0:
                print("✓ LibreOffice 服务正在运行 (端口 2002)")
            else:
                print("✗ LibreOffice 服务未运行")
                print("  启动方法: libreoffice --headless --accept='socket,host=127.0.0.1,port=2002;urp;' &")
        except Exception as e:
            print(f"✗ 检查 LibreOffice 服务时出错: {e}")
    
    return has_win32com or has_uno


def create_test_document():
    """创建测试文档"""
    print_section("创建测试文档")
    
    try:
        from docx import Document
        
        # 创建临时目录
        test_dir = Path(tempfile.gettempdir()) / 'docx_split_test'
        test_dir.mkdir(exist_ok=True)
        
        test_file = test_dir / 'test_document.docx'
        
        # 创建文档
        doc = Document()
        doc.add_heading('测试文档标题', 0)
        
        for i in range(1, 11):
            doc.add_heading(f'第 {i} 页内容', level=1)
            doc.add_paragraph(f'这是第 {i} 页的内容。' * 10)
            doc.add_page_break()
        
        doc.save(test_file)
        print(f"✓ 测试文档已创建: {test_file}")
        
        return str(test_file), str(test_dir)
        
    except ImportError:
        print("✗ python-docx 未安装，无法创建测试文档")
        print("  安装方法: pip install python-docx")
        return None, None
    except Exception as e:
        print(f"✗ 创建测试文档失败: {e}")
        return None, None


def test_split_function(test_file, test_dir):
    """测试拆分功能"""
    print_section("测试 DOCX 拆分功能")
    
    if not test_file or not os.path.exists(test_file):
        print("✗ 测试文件不存在，跳过测试")
        return False
    
    # 创建输出目录
    output_dir = Path(test_dir) / 'output'
    output_dir.mkdir(exist_ok=True)
    
    # 获取脚本路径
    script_dir = Path(__file__).parent / 'server' / 'api' / 'files'
    unified_script = script_dir / 'split_docx_pages_unified.py'
    
    if not unified_script.exists():
        print(f"✗ 脚本不存在: {unified_script}")
        return False
    
    print(f"测试文件: {test_file}")
    print(f"输出目录: {output_dir}")
    print(f"拆分脚本: {unified_script}")
    print(f"每文件页数: 3")
    
    # 执行拆分
    try:
        cmd = [
            sys.executable,
            str(unified_script),
            test_file,
            str(output_dir),
            '3'
        ]
        
        print(f"\n执行命令: {' '.join(cmd)}\n")
        
        result = subprocess.run(
            cmd,
            capture_output=True,
            text=True,
            timeout=60
        )
        
        print("标准输出:")
        print(result.stdout)
        
        if result.stderr:
            print("标准错误:")
            print(result.stderr)
        
        if result.returncode == 0:
            # 检查输出文件
            output_files = list(output_dir.glob('*.docx'))
            print(f"\n✓ 拆分成功！生成了 {len(output_files)} 个文件:")
            for f in output_files:
                size = f.stat().st_size
                print(f"  - {f.name} ({size:,} bytes)")
            return True
        else:
            print(f"\n✗ 拆分失败，退出代码: {result.returncode}")
            return False
            
    except subprocess.TimeoutExpired:
        print("✗ 拆分超时（60秒）")
        return False
    except Exception as e:
        print(f"✗ 执行拆分时出错: {e}")
        import traceback
        traceback.print_exc()
        return False


def cleanup(test_dir):
    """清理测试文件"""
    print_section("清理测试文件")
    
    try:
        import shutil
        if test_dir and os.path.exists(test_dir):
            shutil.rmtree(test_dir)
            print(f"✓ 已删除测试目录: {test_dir}")
    except Exception as e:
        print(f"✗ 清理失败: {e}")


def main():
    """主函数"""
    print_section("跨平台 DOCX 拆分测试")
    
    # 1. 检查依赖
    if not check_dependencies():
        print("\n❌ 缺少必要依赖，无法继续测试")
        print("\n请根据上述提示安装缺失的依赖")
        sys.exit(1)
    
    # 2. 创建测试文档
    test_file, test_dir = create_test_document()
    
    if not test_file:
        print("\n❌ 无法创建测试文档")
        sys.exit(1)
    
    # 3. 测试拆分功能
    success = test_split_function(test_file, test_dir)
    
    # 4. 清理（可选）
    cleanup_choice = input("\n是否清理测试文件？(y/n): ").lower()
    if cleanup_choice == 'y':
        cleanup(test_dir)
    else:
        print(f"\n测试文件保留在: {test_dir}")
    
    # 5. 总结
    print_section("测试总结")
    if success:
        print("✅ 所有测试通过！")
        print("\n系统已正确配置，可以使用跨平台 DOCX 拆分功能")
        sys.exit(0)
    else:
        print("❌ 测试失败")
        print("\n请检查错误信息并修复问题")
        sys.exit(1)


if __name__ == "__main__":
    main()
